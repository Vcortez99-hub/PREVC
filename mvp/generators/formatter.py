import markdown
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from typing import Dict, Any, List, Optional
import os
from datetime import datetime

class DocumentFormatter:
    """Formatador para converter documentação gerada em diferentes formatos"""
    
    def __init__(self):
        self.supported_formats = ['markdown', 'docx', 'html', 'txt']
        
    def format_as_markdown(self, content: str, metadata: Optional[Dict[str, Any]] = None) -> str:
        """Formata conteúdo como Markdown (já está neste formato)"""
        # Adicionar metadata se fornecida
        if metadata:
            header = self._create_metadata_header(metadata)
            return header + "\n\n" + content
        return content
    
    def format_as_html(self, content: str, metadata: Optional[Dict[str, Any]] = None) -> str:
        """Converte Markdown para HTML"""
        # Converter markdown para HTML
        html_content = markdown.markdown(
            content, 
            extensions=['tables', 'toc', 'codehilite']
        )
        
        # Adicionar CSS básico e estrutura HTML
        html_template = f"""
<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Documentação RPA</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            border-left: 4px solid #3498db;
            padding-left: 15px;
            margin-top: 30px;
        }}
        h3 {{
            color: #7f8c8d;
        }}
        ul, ol {{
            padding-left: 20px;
        }}
        li {{
            margin-bottom: 5px;
        }}
        code {{
            background-color: #f8f9fa;
            padding: 2px 4px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
        }}
        .metadata {{
            background-color: #f8f9fa;
            border: 1px solid #e9ecef;
            border-radius: 5px;
            padding: 15px;
            margin-bottom: 20px;
            font-size: 0.9em;
        }}
        .step {{
            background-color: #fff;
            border: 1px solid #dee2e6;
            border-radius: 5px;
            padding: 10px;
            margin: 10px 0;
        }}
    </style>
</head>
<body>
    {self._create_html_metadata(metadata) if metadata else ''}
    {html_content}
    <footer style="margin-top: 50px; padding-top: 20px; border-top: 1px solid #e9ecef; font-size: 0.8em; color: #6c757d;">
        <p>Documentação gerada automaticamente em {datetime.now().strftime('%d/%m/%Y às %H:%M')}</p>
    </footer>
</body>
</html>
"""
        return html_template
    
    def format_as_docx(self, content: str, output_path: str, metadata: Optional[Dict[str, Any]] = None) -> bool:
        """Converte Markdown para documento Word (.docx)"""
        try:
            # Criar novo documento
            doc = Document()
            
            # Configurar estilos
            self._setup_docx_styles(doc)
            
            # Adicionar metadata se fornecida
            if metadata:
                self._add_docx_metadata(doc, metadata)
            
            # Processar conteúdo markdown
            self._parse_markdown_to_docx(doc, content)
            
            # Salvar documento
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Erro ao criar documento Word: {e}")
            return False
    
    def format_as_txt(self, content: str, metadata: Optional[Dict[str, Any]] = None) -> str:
        """Converte para texto puro removendo formatação markdown"""
        # Remover formatação markdown
        text = re.sub(r'#{1,6}\s+', '', content)  # Headers
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
        text = re.sub(r'`(.*?)`', r'\1', text)  # Code
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # Links
        text = re.sub(r'^-\s+', '• ', text, flags=re.MULTILINE)  # Lista não ordenada
        text = re.sub(r'^\d+\.\s+', '', text, flags=re.MULTILINE)  # Lista ordenada
        
        # Adicionar metadata como texto
        if metadata:
            header = self._create_text_metadata(metadata)
            return header + "\n\n" + text
        
        return text
    
    def _create_metadata_header(self, metadata: Dict[str, Any]) -> str:
        """Cria cabeçalho com metadata em formato markdown"""
        header = "---\n"
        header += f"gerado_em: {datetime.now().isoformat()}\n"
        
        if 'process_quality' in metadata:
            pq = metadata['process_quality']
            header += f"qualidade_correlacao: {pq.get('correlation_quality', 0):.2f}\n"
            header += f"acoes_totais: {pq.get('total_actions', 0)}\n"
            header += f"acoes_correlacionadas: {pq.get('correlated_actions', 0)}\n"
        
        if 'document_structure' in metadata:
            ds = metadata['document_structure']
            header += f"secoes_geradas: {ds.get('sections_generated', 0)}\n"
            header += f"passos_identificados: {ds.get('steps_count', 0)}\n"
        
        header += "---\n"
        return header
    
    def _create_html_metadata(self, metadata: Dict[str, Any]) -> str:
        """Cria div com metadata para HTML"""
        html = '<div class="metadata">\n'
        html += '<h3>📊 Informações do Processo</h3>\n'
        
        if 'process_quality' in metadata:
            pq = metadata['process_quality']
            html += f'<p><strong>Qualidade da Correlação:</strong> {pq.get("correlation_quality", 0):.0%}</p>\n'
            html += f'<p><strong>Ações Identificadas:</strong> {pq.get("total_actions", 0)} total, {pq.get("correlated_actions", 0)} correlacionadas</p>\n'
        
        if 'document_structure' in metadata:
            ds = metadata['document_structure']
            html += f'<p><strong>Estrutura:</strong> {ds.get("sections_generated", 0)} seções, {ds.get("steps_count", 0)} passos</p>\n'
        
        html += f'<p><strong>Gerado em:</strong> {datetime.now().strftime("%d/%m/%Y às %H:%M")}</p>\n'
        html += '</div>\n'
        
        return html
    
    def _create_text_metadata(self, metadata: Dict[str, Any]) -> str:
        """Cria metadata para formato texto"""
        text = "=" * 50 + "\n"
        text += "INFORMAÇÕES DO PROCESSO\n"
        text += "=" * 50 + "\n"
        
        if 'process_quality' in metadata:
            pq = metadata['process_quality']
            text += f"Qualidade da Correlação: {pq.get('correlation_quality', 0):.0%}\n"
            text += f"Ações Totais: {pq.get('total_actions', 0)}\n"
            text += f"Ações Correlacionadas: {pq.get('correlated_actions', 0)}\n"
        
        if 'document_structure' in metadata:
            ds = metadata['document_structure']
            text += f"Seções Geradas: {ds.get('sections_generated', 0)}\n"
            text += f"Passos Identificados: {ds.get('steps_count', 0)}\n"
        
        text += f"Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}\n"
        text += "=" * 50
        
        return text
    
    def _setup_docx_styles(self, doc: Document):
        """Configura estilos para documento Word"""
        styles = doc.styles
        
        # Estilo para título principal
        if 'Title Custom' not in [style.name for style in styles]:
            title_style = styles.add_style('Title Custom', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Calibri'
            title_style.font.size = Inches(0.2)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_docx_metadata(self, doc: Document, metadata: Dict[str, Any]):
        """Adiciona metadata ao documento Word"""
        # Adicionar cabeçalho com informações
        header_para = doc.add_paragraph()
        header_para.add_run("📄 DOCUMENTAÇÃO RPA - PROCESSO AUTOMATIZADO").bold = True
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espaço
        
        # Adicionar tabela com metadata
        if 'process_quality' in metadata or 'document_structure' in metadata:
            info_para = doc.add_paragraph()
            info_para.add_run("📊 Informações do Processo").bold = True
            
            info_list = []
            
            if 'process_quality' in metadata:
                pq = metadata['process_quality']
                info_list.extend([
                    f"• Qualidade da Correlação: {pq.get('correlation_quality', 0):.0%}",
                    f"• Ações Identificadas: {pq.get('total_actions', 0)} (total), {pq.get('correlated_actions', 0)} (correlacionadas)",
                ])
            
            if 'document_structure' in metadata:
                ds = metadata['document_structure']
                info_list.extend([
                    f"• Seções Geradas: {ds.get('sections_generated', 0)}",
                    f"• Passos Identificados: {ds.get('steps_count', 0)}",
                ])
            
            info_list.append(f"• Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
            
            for info in info_list:
                doc.add_paragraph(info)
        
        # Adicionar linha separadora
        doc.add_paragraph("_" * 60)
        doc.add_paragraph()
    
    def _parse_markdown_to_docx(self, doc: Document, content: str):
        """Converte conteúdo markdown para elementos do Word"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                doc.add_paragraph()  # Linha em branco
                continue
            
            # Headers
            if line.startswith('# '):
                para = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                para = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                para = doc.add_heading(line[4:], level=3)
            
            # Lista não ordenada
            elif line.startswith('- '):
                para = doc.add_paragraph(line[2:], style='List Bullet')
            
            # Lista ordenada
            elif re.match(r'^\d+\.\s+', line):
                text = re.sub(r'^\d+\.\s+', '', line)
                para = doc.add_paragraph(text, style='List Number')
            
            # Texto normal
            else:
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)
    
    def _add_formatted_text(self, paragraph, text: str):
        """Adiciona texto formatado (bold, italic) ao parágrafo"""
        # Processar formatação markdown no texto
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Code
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
            else:
                # Texto normal
                paragraph.add_run(part)
    
    def export_to_file(self, content: str, format: str, output_path: str, metadata: Optional[Dict[str, Any]] = None) -> bool:
        """Exporta conteúdo para arquivo no formato especificado"""
        try:
            if format.lower() == 'markdown' or format.lower() == 'md':
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(self.format_as_markdown(content, metadata))
                return True
            
            elif format.lower() == 'html':
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(self.format_as_html(content, metadata))
                return True
            
            elif format.lower() == 'txt':
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(self.format_as_txt(content, metadata))
                return True
            
            elif format.lower() == 'docx':
                return self.format_as_docx(content, output_path, metadata)
            
            else:
                print(f"Formato não suportado: {format}")
                return False
                
        except Exception as e:
            print(f"Erro ao exportar arquivo: {e}")
            return False
    
    def get_supported_formats(self) -> List[str]:
        """Retorna lista de formatos suportados"""
        return self.supported_formats.copy()
    
    def validate_output_path(self, output_path: str, format: str) -> bool:
        """Valida se o caminho de saída é válido para o formato"""
        try:
            # Verificar se diretório existe
            directory = os.path.dirname(output_path)
            if directory and not os.path.exists(directory):
                os.makedirs(directory, exist_ok=True)
            
            # Verificar extensão
            expected_extensions = {
                'markdown': ['.md', '.markdown'],
                'html': ['.html', '.htm'],
                'txt': ['.txt'],
                'docx': ['.docx']
            }
            
            if format.lower() in expected_extensions:
                file_ext = os.path.splitext(output_path)[1].lower()
                return file_ext in expected_extensions[format.lower()]
            
            return False
            
        except Exception:
            return False